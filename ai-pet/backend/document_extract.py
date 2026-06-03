from __future__ import annotations

import os
import sys
from pathlib import Path

MAX_CHARS = int(os.getenv("FILE_CHAT_MAX_BYTES", "80000"))
MAX_PDF_PAGES = int(os.getenv("FILE_CHAT_MAX_PDF_PAGES", "40"))


def _truncate(text: str) -> str:
    cleaned = text.replace("\r\n", "\n").strip()
    if not cleaned:
        return ""
    if len(cleaned) <= MAX_CHARS:
        return cleaned
    return cleaned[:MAX_CHARS] + f"\n\n…（已截断，仅展示前 {MAX_CHARS} 字）"


def extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "（未安装 pypdf，无法解析 PDF。请运行: pip install -r backend/requirements.txt）"

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        return f"（PDF 打开失败：{exc}）"

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:
            return "（PDF 已加密，请先解密或导出为无密码版本。）"

    if not reader.pages:
        return "（PDF 无页面）"

    parts: list[str] = []
    for index, page in enumerate(reader.pages[:MAX_PDF_PAGES]):
        try:
            chunk = page.extract_text() or ""
        except Exception as exc:
            chunk = f"（第 {index + 1} 页提取失败：{exc}）"
        if chunk.strip():
            parts.append(chunk.strip())

    if len(reader.pages) > MAX_PDF_PAGES:
        parts.append(f"（仅载入前 {MAX_PDF_PAGES} 页，共 {len(reader.pages)} 页）")

    text = "\n\n".join(parts)
    if not text.strip():
        return (
            "（未能从 PDF 提取文字，可能是扫描件/图片 PDF。"
            "请先 OCR 或复制文字到 .txt 后再拖入。）"
        )
    return _truncate(text)


def extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return "（未安装 python-docx，无法解析 Word。请运行: pip install -r backend/requirements.txt）"

    try:
        document = Document(str(path))
    except Exception as exc:
        return f"（Word 文档打开失败：{exc}）"

    lines: list[str] = []
    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        if line:
            lines.append(line)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append("\t".join(cells))

    if not lines:
        return "（Word 文档中没有可提取的文字）"
    return _truncate("\n".join(lines))


def extract_doc_win32(path: Path) -> str | None:
    if sys.platform != "win32":
        return None
    try:
        import win32com.client  # type: ignore[import-untyped]
    except ImportError:
        return None

    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(path.resolve()), ReadOnly=True)
        text = str(doc.Content.Text or "")
        return _truncate(text) if text.strip() else None
    except Exception:
        return None
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def extract_doc(path: Path) -> str:
    extracted = extract_doc_win32(path)
    if extracted:
        return extracted
    if sys.platform == "win32":
        return (
            "（无法读取 .doc 老格式：请安装 Microsoft Word，"
            "或在 Word 中「另存为」.docx 后再拖入。）"
        )
    return "（.doc 老格式请在 Windows 上另存为 .docx 后再拖入。）"


def extract_document(path: Path) -> str | None:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".doc":
        return extract_doc(path)
    return None
